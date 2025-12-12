from typing import List
from io import BytesIO
from docx import Document

from .base import DocumentProcessor, ExtractedText


class WordProcessor(DocumentProcessor):
    """Processor for Word documents using python-docx"""
    
    SUPPORTED_EXTENSIONS = ['docx', 'doc']
    
    async def extract_text(self, content: bytes, filename: str) -> List[ExtractedText]:
        """
        Extract text from a Word document.
        
        Args:
            content: Raw bytes of the Word file
            filename: Original filename
            
        Returns:
            List of ExtractedText objects
        """
        try:
            doc_file = BytesIO(content)
            doc = Document(doc_file)
            
            # Extract all text from paragraphs
            paragraphs_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs_text.append(para.text.strip())
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = "\n\n".join(paragraphs_text)
            if tables_text:
                all_text += "\n\n--- Tablas ---\n" + "\n".join(tables_text)
            
            if not all_text.strip():
                return []
            
            # Return as single extracted text (Word docs don't have clear page breaks)
            return [ExtractedText(
                text=all_text,
                page_number=1,
                metadata={
                    "source_file": filename,
                    "file_type": "word",
                    "paragraphs_count": len(paragraphs_text),
                    "tables_count": len(doc.tables)
                }
            )]
            
        except Exception as e:
            raise Exception(f"Error extracting text from Word document: {str(e)}")
    
    def supports_format(self, filename: str) -> bool:
        """Check if this processor supports Word files"""
        ext = self.get_extension(filename)
        return ext in self.SUPPORTED_EXTENSIONS
